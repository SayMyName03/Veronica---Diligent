"""
Document processor for handling PDF and DOCX files
"""
import os
from typing import List
from pypdf import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
from config import Config


class DocumentProcessor:
    """Process PDF and DOCX documents"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            length_function=len,
        )
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    def process_document(self, file_path: str) -> List[LangchainDocument]:
        """
        Process a document and return chunks
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of LangChain Document objects
        """
        # Extract text based on file extension
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            text = self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        # Create metadata
        metadata = {
            "source": os.path.basename(file_path),
            "file_type": file_ext
        }
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create LangChain documents
        documents = [
            LangchainDocument(page_content=chunk, metadata=metadata)
            for chunk in chunks
        ]
        
        return documents
    
    def save_uploaded_file(self, uploaded_file, filename: str) -> str:
        """
        Save uploaded file to disk
        
        Args:
            uploaded_file: File object from Streamlit
            filename: Name of the file
            
        Returns:
            Path to saved file
        """
        # Create upload directory if it doesn't exist
        os.makedirs(Config.UPLOAD_DIR, exist_ok=True)
        
        # Save file
        file_path = os.path.join(Config.UPLOAD_DIR, filename)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        return file_path
